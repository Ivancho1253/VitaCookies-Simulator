from __future__ import annotations

from datetime import datetime
from io import BytesIO
from math import isinf
from typing import Any


MODEL_CARDS = {
    "digital": {
        "title": "Analisis 1 - Flujo real del formulario digital",
        "system": "Evento de testeo sensorial ya realizado, con respuestas registradas en el formulario digital.",
        "objective": "Medir carga real, pico por minuto y funcionamiento operativo del formulario.",
        "entities": "Comensales, formulario digital, servidor/capacidad concurrente.",
        "state": "Respuestas por minuto, acumulado, utilizacion observada y minutos saturados.",
        "events": "Envio real de respuestas del formulario.",
        "parameters": "Respuestas registradas, duracion observada y capacidad aceptable por minuto.",
        "inputs": "Datos reales del formulario y capacidad operativa editable.",
        "outputs": "Respuestas, duracion, pico por minuto, utilizacion maxima y decision sugerida.",
        "assumptions": "No se simulan llegadas; se usan las marcas horarias reales del Excel.",
        "restrictions": "No mide fallas de red no registradas en el formulario.",
        "scope": "Apoya la lectura posterior del funcionamiento digital del testeo.",
    },
    "stock": {
        "title": "Analisis 2 - Stock real de galletitas",
        "system": "Inventario de galletitas producidas para el testeo sensorial ya realizado.",
        "objective": "Medir consumo real, sobrantes y cantidad sugerida para un proximo testeo similar.",
        "entities": "Comensales, porciones, desperdicio, demanda.",
        "state": "Galletitas producidas, consumidas y sobrantes.",
        "events": "Produccion, consumo durante el testeo y sobrante final.",
        "parameters": "Galletitas producidas, sobrantes aproximadas y respuestas del formulario.",
        "inputs": "Unidades reales del testeo.",
        "outputs": "Consumo observado, sobrante observado y produccion sugerida.",
        "assumptions": "Las sobrantes son aproximadas segun el dato informado por el equipo.",
        "restrictions": "No separa sobrante por descarte, rotura o no consumo si no fue medido.",
        "scope": "Apoya la decision de cantidad a producir en un proximo testeo similar.",
    },
    "viability": {
        "title": "Analisis 3 - Aceptabilidad y escala productiva",
        "system": "Produccion a mayor escala usando la aceptabilidad real obtenida en el testeo sensorial.",
        "objective": "Proyectar costo, ingresos y ganancia para una cantidad mayor de galletitas.",
        "entities": "Galletitas producidas, unidades aceptadas estimadas, consumidores, costo de 50 unidades, precio e ingresos.",
        "state": "Aceptabilidad, costo unitario estimado, produccion objetivo, ingresos y ganancia estimada.",
        "events": "Carga del costo real de producir 50 galletitas y proyeccion a escala.",
        "parameters": "Costo de producir 50 galletitas, precio unitario y cantidad a producir.",
        "inputs": "Aceptabilidad real del testeo y datos economicos editables.",
        "outputs": "Aceptabilidad positiva, unidades aceptadas estimadas, costo unitario, precio de equilibrio, ingresos y ganancia estimada.",
        "assumptions": "La aceptabilidad real de 41 respuestas positivas sobre 42 se usa como tasa base para escalar.",
        "restrictions": "No reemplaza un estudio de mercado ni costos industriales reales.",
        "scope": "Apoya una decision posterior sobre escala de produccion, precio y ajuste de receta.",
    },
}


def _fmt(value: Any) -> str:
    if value is None:
        return "No ejecutado"
    if isinstance(value, tuple):
        return f"{value[0]:,.2f} a {value[1]:,.2f}"
    if isinstance(value, float) and isinf(value):
        return "No alcanzable"
    if isinstance(value, float):
        return f"{value:,.2f}"
    return str(value)


def _model_section(key: str, result: dict | None) -> str:
    card = MODEL_CARDS[key]
    lines = [
        f"## {card['title']}",
        "",
        "### Modelado formal",
        f"- **Sistema:** {card['system']}",
        f"- **Objetivo:** {card['objective']}",
        f"- **Entidades:** {card['entities']}",
        f"- **Variables de estado:** {card['state']}",
        f"- **Eventos:** {card['events']}",
        f"- **Parametros:** {card['parameters']}",
        f"- **Variables de entrada:** {card['inputs']}",
        f"- **Variables de salida:** {card['outputs']}",
        f"- **Supuestos:** {card['assumptions']}",
        f"- **Restricciones:** {card['restrictions']}",
        f"- **Alcance:** {card['scope']}",
        "",
    ]
    if not result:
        lines.append("El simulador no fue ejecutado en esta sesion.")
        return "\n".join(lines)

    metrics = result["metrics"]
    lines.extend(["### Resultado, interpretacion y decision", ""])
    lines.append(f"- **Resultado:** {metrics.get('resultado', 'No disponible')}")
    lines.append(f"- **Interpretacion:** {metrics.get('interpretacion', 'No disponible')}")
    lines.append(f"- **Recomendacion:** {metrics.get('recomendacion', 'No disponible')}")
    lines.append(f"- **Decision sugerida:** {metrics.get('decision', 'No disponible')}")
    lines.append("")
    lines.append("### Indicadores principales")
    for key_metric, value in metrics.items():
        if key_metric in {"resultado", "interpretacion", "recomendacion", "decision"}:
            continue
        lines.append(f"- **{key_metric.replace('_', ' ').capitalize()}:** {_fmt(value)}")
    return "\n".join(lines)


def generate_markdown_report(results: dict[str, dict | None], verification: dict[str, dict] | None = None) -> str:
    date = datetime.now().strftime("%d/%m/%Y %H:%M")
    verification = verification or {}
    verification_lines = []
    for model, checks in verification.items():
        verification_lines.append(f"### {model}")
        for check, ok in checks.items():
            verification_lines.append(f"- {check.replace('_', ' ')}: {'cumple' if ok else 'revisar'}")
    verification_text = "\n".join(verification_lines) if verification_lines else "La verificacion se completa al cargar los datos post-testeo."

    return f"""# Informe academico post-testeo - VitaCookies

**Fecha:** {date}

**Materia:** Modelos y Simulacion  
**Proyecto:** Evaluacion integradora intercatedra entre Ingenieria en Sistemas y Nutricion  
**Producto:** Galletita vegetal sustentable con avena, lentejas, manzana y zanahoria.

## Portada

El presente informe documenta una herramienta post-testeo desarrollada para leer los datos reales obtenidos en la evaluacion sensorial de VitaCookies. La herramienta resume funcionamiento digital, stock observado, aceptacion sensorial y viabilidad productiva/comercial.

## Descripcion general

El objetivo ya no es predecir el evento, sino transformar las mediciones obtenidas en indicadores claros y recomendaciones concretas para el equipo de Nutricion.

{_model_section("digital", results.get("digital"))}

{_model_section("stock", results.get("stock"))}

{_model_section("viability", results.get("viability"))}

## Datos observados

La version post-testeo reemplaza corridas aleatorias, escenarios y probabilidades simuladas por valores reales: respuestas del formulario, horarios de carga, galletitas producidas, galletitas sobrantes, aceptacion y puntajes descriptivos.

## Verificacion

{verification_text}

## Validacion del modelo

La validacion queda incorporada porque los calculos usan los datos reales del evento: cantidad de respuestas, tiempo observado de carga del formulario, momentos de mayor concurrencia, galletitas consumidas, sobrantes, aceptacion sensorial e indicadores descriptivos.

## Analisis para toma de decisiones

- Si el pico digital observado se acerca a la capacidad definida, se recomienda escalonar respuestas en futuros testeos.
- Si el sobrante observado es bajo, la produccion estuvo bien dimensionada.
- Si la aceptabilidad es alta pero la ganancia estimada es baja, se recomienda revisar precio o costo de producir 50 unidades.
- Si textura aparece como atributo debil, se recomienda mejorar crocancia sin perder sabor.

## Limitaciones

- Los indicadores resumen comportamientos individuales en metricas agregadas.
- La capacidad del formulario es estimada si no se dispone de medicion tecnica real.
- Los costos comerciales deben reemplazarse por datos reales si el producto escala.
- Las sobrantes son aproximadas segun el dato informado por el equipo.

## Mejoras futuras

- Automatizar la importacion directa del Excel del formulario.
- Guardar historiales de testeos.
- Comparar automaticamente plan previo vs resultado post-testeo.
- Exportar graficos al informe final.

## Guia para defensa oral

1. **Formulario digital:** explicar el volumen real de respuestas y el pico observado.
2. **Stock:** explicar que 50 galletitas producidas y 5 sobrantes aproximadas muestran un stock bien dimensionado.
3. **Aceptacion:** destacar la aceptacion positiva y los atributos descriptivos.
4. **Viabilidad:** mostrar como precio y costos transforman el consumo real en resultado economico.
5. **Conclusion:** la herramienta transforma datos reales del testeo en decisiones concretas para Nutricion.
"""


def generate_docx_report(markdown_text: str) -> bytes:
    from docx import Document

    document = Document()
    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            document.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            document.add_heading(line[4:], level=2)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        elif line[0:2].isdigit() and ". " in line[:4]:
            document.add_paragraph(line, style="List Number")
        else:
            document.add_paragraph(line.replace("**", ""))
    output = BytesIO()
    document.save(output)
    return output.getvalue()
